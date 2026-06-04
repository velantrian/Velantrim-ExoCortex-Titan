"""
📝 DOCX Generator v1.0 — Microsoft Word через python-docx
============================================================
Профессиональные Word-документы из GenerationSpec.

Особенности:
- Стили под темы (clean/scientific/business/dark/velantrim)
- Multi-level headings (h1-h6)
- Таблицы с alternating row colors
- Callout-блоки через single-cell tables с border-color
- Code blocks с моноширинным шрифтом и серым фоном
- Header/footer с pagination и метаданными
- Поддержка кириллицы
- Tracked changes готовность (если нужно — можно расширить)
"""

import logging
import time

from .base import (
    CalloutBlock,
    CodeBlock,
    DividerBlock,
    FactBlock,
    FileGenerator,
    GenerationResult,
    GenerationSpec,
    HeadingBlock,
    ImageBlock,
    ListBlock,
    ParagraphBlock,
    QuoteBlock,
    TableBlock,
    get_theme,
)

logger = logging.getLogger("velantrim.generators.docx")


def _check_available(module_name: str) -> bool:
    import importlib.util
    return importlib.util.find_spec(module_name) is not None


PYTHON_DOCX_AVAILABLE = _check_available("docx")


class DOCXGenerator(FileGenerator):
    format_name = "docx"

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def generate(
        self,
        spec: GenerationSpec,
        output_path: str,
    ) -> GenerationResult:
        result = GenerationResult(output_path=output_path, format="docx")
        start = time.time()

        if not PYTHON_DOCX_AVAILABLE:
            result.error = "python-docx не установлен. pip install python-docx"
            return result

        self._ensure_output_dir(output_path)
        theme = get_theme(spec.theme)

        try:
            self._render(spec, output_path, theme)
            result.method = "python-docx"
            result.block_count = len(spec.blocks)
            result.file_size_bytes = self._file_size(output_path)
            result.metadata = self._build_provenance(output_path)
        except Exception as exc:
            logger.error(f"DOCX generation error: {exc}")
            result.error = str(exc)

        result.generation_time_ms = (time.time() - start) * 1000
        return result

    # ─── Render implementation ────────────────────────────────────────────────

    def _render(self, spec: GenerationSpec, output_path: str, theme) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()

        # ─── Метаданные документа ───
        doc.core_properties.title = spec.metadata.title
        doc.core_properties.author = spec.metadata.author
        doc.core_properties.subject = spec.metadata.subject
        if spec.metadata.keywords:
            doc.core_properties.keywords = ", ".join(spec.metadata.keywords)

        # ─── Стили под тему ───
        styles = doc.styles
        self._setup_styles(styles, theme)

        # ─── Footer ───
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Velantrim ExoCortex • {spec.metadata.author}"
        )
        footer_run.font.size = Pt(theme.size_xs)
        footer_run.font.color.rgb = self._hex_to_rgb(theme.text_muted)

        # ─── Render блоков ───
        for block in spec.blocks:
            if isinstance(block, HeadingBlock):
                self._add_heading(doc, block, theme)
            elif isinstance(block, ParagraphBlock):
                self._add_paragraph(doc, block, theme)
            elif isinstance(block, ListBlock):
                self._add_list(doc, block)
            elif isinstance(block, TableBlock):
                self._add_table(doc, block, theme)
            elif isinstance(block, CodeBlock):
                self._add_code(doc, block, theme)
            elif isinstance(block, ImageBlock):
                self._add_image(doc, block)
            elif isinstance(block, CalloutBlock):
                self._add_callout(doc, block, theme)
            elif isinstance(block, QuoteBlock):
                self._add_quote(doc, block, theme)
            elif isinstance(block, DividerBlock):
                self._add_divider(doc, theme)
            elif isinstance(block, FactBlock):
                self._add_fact(doc, block, theme)

        doc.save(output_path)

    # ─── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _hex_to_rgb(hex_color: str):
        from docx.shared import RGBColor
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )

    def _setup_styles(self, styles, theme) -> None:
        """Настройка стилей под тему."""
        from docx.shared import Pt

        # Базовый стиль
        normal = styles["Normal"]
        normal.font.name = theme.font_body
        normal.font.size = Pt(theme.size_md)

    def _add_heading(self, doc, block: HeadingBlock, theme) -> None:
        from docx.shared import Pt
        level = min(block.level, 9)
        h = doc.add_heading(block.text, level=level)
        # Цвет заголовка
        for run in h.runs:
            run.font.color.rgb = self._hex_to_rgb(theme.primary)
            run.font.name = theme.font_heading
            size_map = {
                1: theme.size_3xl, 2: theme.size_2xl,
                3: theme.size_xl, 4: theme.size_lg,
            }
            run.font.size = Pt(size_map.get(level, theme.size_md))

    def _add_paragraph(self, doc, block: ParagraphBlock, theme) -> None:
        from docx.shared import Pt
        p = doc.add_paragraph()
        run = p.add_run(block.text)
        run.font.size = Pt(theme.size_md)
        run.font.color.rgb = self._hex_to_rgb(theme.text)
        if block.style == "bold":
            run.bold = True
        elif block.style == "italic":
            run.italic = True
        elif block.style == "callout":
            run.font.color.rgb = self._hex_to_rgb(theme.text_muted)
            run.italic = True

    def _add_list(self, doc, block: ListBlock) -> None:
        style = "List Number" if block.ordered else "List Bullet"
        for item in block.items:
            doc.add_paragraph(item, style=style)

    def _add_table(self, doc, block: TableBlock, theme) -> None:
        from docx.shared import Pt

        # Создаём таблицу с headers + rows
        ncols = len(block.headers) if block.headers else (
            len(block.rows[0]) if block.rows else 1
        )
        table = doc.add_table(rows=1, cols=ncols)
        table.style = "Light Grid Accent 1"

        # Header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(block.headers):
            cell = hdr_cells[i]
            cell.text = str(header)
            # Стилизуем header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(theme.size_md)
                    run.font.color.rgb = self._hex_to_rgb(theme.background)
            # Фон header
            self._set_cell_bg(cell, theme.primary)

        # Body rows
        for row_data in block.rows:
            row = table.add_row()
            for i, cell_val in enumerate(row_data):
                if i < ncols:
                    cell = row.cells[i]
                    cell.text = str(cell_val) if cell_val is not None else ""
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(theme.size_sm)

        if block.caption:
            cap = doc.add_paragraph(block.caption)
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(theme.size_sm)
                run.font.color.rgb = self._hex_to_rgb(theme.text_muted)

    @staticmethod
    def _set_cell_bg(cell, hex_color: str) -> None:
        """Установить фон ячейки таблицы."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    def _add_code(self, doc, block: CodeBlock, theme) -> None:
        from docx.shared import Pt
        # Code в single-cell таблице с серым фоном
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(block.code)
        run.font.name = theme.font_mono
        run.font.size = Pt(theme.size_sm)
        run.font.color.rgb = self._hex_to_rgb(theme.text)
        self._set_cell_bg(cell, theme.surface)

        if block.caption:
            cap = doc.add_paragraph(block.caption)
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(theme.size_xs)
                run.font.color.rgb = self._hex_to_rgb(theme.text_muted)

    def _add_image(self, doc, block: ImageBlock) -> None:
        import os

        from docx.shared import Inches
        if not os.path.exists(block.path):
            doc.add_paragraph(f"[Изображение не найдено: {block.path}]")
            return
        try:
            width = Inches(block.width / 96) if block.width else Inches(5)
            doc.add_picture(block.path, width=width)
            if block.caption:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                cap = doc.add_paragraph(block.caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.italic = True
        except Exception as exc:
            doc.add_paragraph(f"[Ошибка изображения: {exc}]")

    def _add_callout(self, doc, block: CalloutBlock, theme) -> None:
        from docx.shared import Pt
        # Single-cell table с цветной границей слева
        type_config = {
            "info": ("ℹ️", theme.primary),
            "success": ("✅", theme.success),
            "warning": ("⚠️", theme.warning),
            "danger": ("🚨", theme.danger),
        }
        emoji, color = type_config.get(block.callout_type, type_config["info"])

        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = ""
        self._set_cell_bg(cell, theme.surface)

        # Заголовок callout
        para = cell.paragraphs[0]
        if block.title:
            run = para.add_run(f"{emoji} {block.title}")
            run.bold = True
            run.font.size = Pt(theme.size_md)
            run.font.color.rgb = self._hex_to_rgb(color)
            para.add_run("\n")

        # Текст
        text_run = para.add_run(block.text)
        text_run.font.size = Pt(theme.size_md)
        text_run.font.color.rgb = self._hex_to_rgb(theme.text)

    def _add_quote(self, doc, block: QuoteBlock, theme) -> None:
        from docx.shared import Pt
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(20)
        run = p.add_run(f'"{block.text}"')
        run.italic = True
        run.font.size = Pt(theme.size_md)
        run.font.color.rgb = self._hex_to_rgb(theme.text_muted)
        if block.author:
            p.add_run(f"\n— {block.author}").italic = True

    def _add_divider(self, doc, theme) -> None:
        """Горизонтальная линия — через нижнюю границу пустого параграфа."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), theme.border)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_fact(self, doc, block: FactBlock, theme) -> None:
        """Velantrim-специфичный рендеринг факта."""
        from docx.shared import Pt
        state_colors = {
            "Validated": theme.success,
            "Supported": theme.success,
            "Hypothesized": theme.warning,
            "Observed": theme.primary,
            "Contradicted": theme.danger,
            "Collapsed": theme.danger,
            "Deprecated": theme.danger,
            "ImmutableCore": theme.success,
        }
        state_colors.get(block.epistemic_state, theme.primary)
        conf_label = self._confidence_to_label(block.confidence)

        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = ""
        self._set_cell_bg(cell, theme.surface)

        para = cell.paragraphs[0]
        # Claim
        claim_run = para.add_run(block.claim)
        claim_run.bold = True
        claim_run.font.size = Pt(theme.size_md)
        claim_run.font.color.rgb = self._hex_to_rgb(theme.text)

        # Метаданные
        meta_para = cell.add_paragraph()
        meta_text = (
            f"ID: {block.fact_id}  |  "
            f"Состояние: {block.epistemic_state}  |  "
            f"Уверенность: {block.confidence:.3f} {conf_label}  |  "
            f"Источник: {block.source}"
        )
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.size = Pt(theme.size_xs)
        meta_run.font.color.rgb = self._hex_to_rgb(theme.text_muted)
