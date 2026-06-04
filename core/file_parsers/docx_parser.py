"""
📝 DOCX Parser v2.0 — Microsoft Word документы
================================================
Каскад: Unstructured → python-docx

CHANGELOG v2.0:
- 🆕 _enrich_with_essence унаследован из base (был дубликат)
- 🆕 Tracked changes detection
- 🆕 Comments extraction
- 🆕 word_count, page_count
"""

import logging
import time

from .base import FileParser, ParseResult

logger = logging.getLogger("velantrim.parsers.docx")


def _check_available(module_name: str) -> bool:
    import importlib.util
    return importlib.util.find_spec(module_name) is not None


PYTHON_DOCX_AVAILABLE  = _check_available("docx")
UNSTRUCTURED_AVAILABLE = _check_available("unstructured")


class DOCXParser(FileParser):
    file_type = "docx"

    def supported_formats(self) -> list:
        return [".docx", ".docm"]

    def parse(self, file_path: str) -> ParseResult:
        early = self._check_file(file_path)
        if early is not None:
            return early

        result = ParseResult(file_path=file_path, file_type="docx")
        result.file_size_bytes = self._get_file_size(file_path)
        start = time.time()

        try:
            if UNSTRUCTURED_AVAILABLE:
                text, structured = self._parse_unstructured(file_path)
                result.extraction_method = "Unstructured"
            elif PYTHON_DOCX_AVAILABLE:
                text, structured = self._parse_python_docx(file_path)
                result.extraction_method = "python-docx"
            else:
                result.error = (
                    "Ни один DOCX-парсер не доступен. "
                    "pip install python-docx ИЛИ pip install unstructured[docx]"
                )
                return result

            result.extracted_text = text
            result.structured_data = structured
            result.word_count = len(text.split()) if text else 0
            result.provenance = self._build_provenance(
                file_path, result.extraction_method
            )

            if text.strip():
                result = self._enrich_with_essence(result)

        except Exception as exc:
            logger.error(f"DOCXParser error: {exc}")
            result.error = str(exc)

        result.parse_time_ms = (time.time() - start) * 1000
        return result

    def _parse_unstructured(self, file_path: str) -> tuple[str, dict]:
        from unstructured.partition.docx import partition_docx
        elements = partition_docx(file_path)
        text = "\n\n".join(str(el) for el in elements)
        return text, {
            "elements": [
                {"type": type(el).__name__, "text": str(el)}
                for el in elements
            ],
        }

    def _parse_python_docx(self, file_path: str) -> tuple[str, dict]:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables = []
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(rows)

        # v2.0: comments extraction
        comments = []
        try:
            # python-docx не имеет прямого API для comments,
            # но они доступны через XML
            from docx.oxml.ns import qn
            comments_part = doc.part.related_parts.get("comments")
            if comments_part:
                comments_xml = comments_part.element
                for c in comments_xml.findall(qn("w:comment")):
                    author = c.get(qn("w:author"), "")
                    text = "".join(t.text or "" for t in c.iter(qn("w:t")))
                    comments.append({"author": author, "text": text})
        except Exception:
            pass

        # v2.0: tracked changes detection
        has_tracked_changes = False
        try:
            from docx.oxml.ns import qn
            body = doc.element.body
            if body.find(qn("w:ins")) is not None or body.find(qn("w:del")) is not None:
                has_tracked_changes = True
        except Exception:
            pass

        return "\n\n".join(paragraphs), {
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "comments": comments,
            "has_tracked_changes": has_tracked_changes,
        }
